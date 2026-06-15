import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, hex_color):
    """Establece el color de fondo de una celda usando manipulación XML de Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_mapping_document():
    # Inicializar documento
    doc = docx.Document()
    
    # Configurar márgenes a 1 pulgada (2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos de fuente por defecto (Calibri/Calibri Light o Arial)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Título Principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title.add_run("Mapeo de Integración IFTMIN (XML) ➔ UNIGIS (BD)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Azul Marino
    title.paragraph_format.space_after = Pt(6)

    # Subtítulo
    subtitle = doc.add_paragraph()
    run_sub = subtitle.add_run("Validación, Correcciones de Concepto y Completitud de Campos")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    subtitle.paragraph_format.space_after = Pt(24)

    # Párrafo introductorio
    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "Este documento presenta el análisis y mapeo validado para la integración de los mensajes XML de "
        "instrucción de transporte "
    )
    bold_iftmin = p_intro.add_run("IFTMIN (EDIFACT / Portic)")
    bold_iftmin.bold = True
    p_intro.add_run(" hacia el sistema de base de datos de ")
    bold_unigis = p_intro.add_run("UNIGIS TMS (Europastry)")
    bold_unigis.bold = True
    p_intro.add_run(
        ". Las modificaciones resuelven inconsistencias críticas de negocio encontradas en el borrador original "
        "y completan campos obligatorios para el flujo terrestre/marítimo."
    )
    p_intro.paragraph_format.space_after = Pt(18)

    # Heading 1: Tabla de Mapeo
    h1_table = doc.add_paragraph()
    run_h1 = h1_table.add_run("1. Tabla de Mapeo Validada y Completada")
    run_h1.font.name = 'Arial'
    run_h1.font.size = Pt(15)
    run_h1.font.bold = True
    run_h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_table.paragraph_format.space_before = Pt(12)
    h1_table.paragraph_format.space_after = Pt(8)

    # Datos para la tabla
    headers = [
        "Campo XML (IFTMIN)", 
        "Campo BD Original", 
        "Campo BD Propuesto", 
        "Estado", 
        "Descripción y Notas"
    ]
    
    rows_data = [
        ("tred_document.message.number", "Viaje.ReferenciaExteterna", "Viaje.ReferenciaExterna", "Corregido", "Corrección de error tipográfico (ReferenciaExteterna -> ReferenciaExterna). ID externo del viaje."),
        ("tred_date.time.period (137)", "Viaje.Fecha", "Viaje.Fecha o Viaje.FechaCreacion", "Correcto", "Calificador 137 representa la fecha de emisión del mensaje/documento."),
        ("tred_id.of.the.means.of.transport", "Vehiculo.Dominio", "Vehiculo.Dominio", "Correcto", "Matrícula del camión (patente) o ID del buque (TDT)."),
        ("tred_carrier.name", "Empresa.RazonSocial", "Transporte.RazonSocial", "Corregido", "En UNIGIS, Empresa es el dador de carga (Tenant). El carrier se modela en la entidad Transporte (o Transportista)."),
        ("tred_date.time.period (200)", "Viaje.FechaInicioPlan", "Viaje.FechaInicioPlan", "Correcto", "Calificador 200 representa la fecha estimada de salida/inicio plan."),
        ("tred_date.time.period (2)", "Viaje.FechaFinPlan", "Viaje.FechaFinPlan", "Correcto", "Calificador 2 representa la fecha estimada de entrega/fin de viaje."),
        ("tred_party.id.identification", "Transporte.CUIT", "Transporte.CUIT (o .ReferenciaExterna)", "Correcto", "Identificación fiscal del transportista (si la parte es CA)."),
        ("tred_name.and.address.line", "Empresa.RazonSocial", "Sucursal.Nombre", "Corregido", "Al estar dentro del bloque NAD que describe el destino/origen, representa el nombre de la sucursal/parada, no de la Empresa."),
        ("tred_street.and.number...", "Sucursal.Direccion", "Sucursal.Direccion", "Correcto", "Dirección física de la parada."),
        ("tred_city.name", "Sucursal.Localidad", "Sucursal.Localidad", "Correcto", "Localidad/Ciudad del punto de parada."),
        ("tred_post.code", "Sucursal.CodigoPostal", "Sucursal.CodigoPostal", "Correcto", "Código Postal."),
        ("trcd_reference (BN)", "[Vacío]", "Viaje.ReferenciaAdicional\n(o Viaje.Z_BookingNumber)", "Completado", "BN = Booking Number. Si Viaje.ReferenciaExterna ya se usa para el ID de mensaje, se guarda en campos adicionales."),
        ("tred_goods.item.number", "ParadaItem.Orden", "ParadaItem.Orden", "Correcto", "Índice o renglón de ítem en la parada."),
        ("tred_number.of.packages", "ParadaItem.Cantidad", "ParadaItem.Cantidad (o .Bultos)", "Correcto", "Cantidad de bultos declarados para el ítem."),
        ("tred_free.text (AAA)", "ParadaItem.Descripcion", "ParadaItem.Descripcion", "Correcto", "Calificador AAA representa la descripción de la mercadería."),
        ("tred_type.of.packages.id", "ParadaItem.TipoITem", "ParadaItem.TipoItem (o .Z_TipoBulto)", "Corregido", "Corrección de mayúsculas (TipoITem -> TipoItem). Tipo de embalaje (Caja, Palet, etc.)."),
        ("tred_equipment.id.number", "ParadaItem_Dyn.Z_PromocionCampana", "Parada_Dyn.Z_Contenedor\n(o UnidadContenedora.RefExt)", "Corregido", "Grave: Este campo lleva la sigla/número del contenedor (ej. MSKU1234567). No tiene relación con una promoción/campaña."),
        ("anxe_sender.identification", "Parada.Idcliente", "Parada.IdCliente\n(o Viaje.Z_TerminalID)", "Corregido", "Corrección de mayúsculas (Idcliente -> IdCliente). Representa la terminal/depósito origen del mensaje."),
        ("tred_transport.stage.qualifier", "Viaje_Dyn.Etapa", "Viaje_Dyn.Etapa", "Correcto", "Define la etapa del transporte (ej. principal, previo)."),
        ("tred_seal.number", "Viaje_Dyn.IdPrecinto", "Parada_Dyn.Z_Precinto\n(o Viaje_Dyn.IdPrecinto)", "Corregido", "El precinto se asocia físicamente al contenedor. Dado que 1 Contenedor = 1 Parada, se recomienda a nivel parada.")
    ]

    # Crear tabla: filas = len(data) + 1 (para header), cols = 5
    table = doc.add_table(rows=len(rows_data)+1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Anchos de columna propuestos (suman 6.5 pulgadas)
    col_widths = [Inches(1.5), Inches(1.2), Inches(1.3), Inches(0.8), Inches(1.7)]

    # Dar formato al Header
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        # Aplicar color de fondo oscuro azul
        set_cell_background(hdr_cells[i], "1F4E78")
        # Formato de texto del Header
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # Blanco
            run.font.size = Pt(9.5)

    # Rellenar datos
    for r_idx, row_val in enumerate(rows_data):
        row_cells = table.rows[r_idx+1].cells
        for c_idx, val in enumerate(row_val):
            row_cells[c_idx].text = val
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            # Formatear el tamaño de fuente en las celdas
            for run in p.runs:
                run.font.size = Pt(8.5)
                
            # Formatear columna Estado con colores específicos
            if c_idx == 3: # Columna Estado
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    if val == "Corregido":
                        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Rojo oscuro
                    elif val == "Completado":
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Verde oscuro
                    else:
                        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Gris
            
            # Alinear a la izquierda otras columnas
            elif c_idx == 4:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Aplicar anchos de columna a todas las celdas
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 2: Explicación de Correcciones
    h1_corr = doc.add_paragraph()
    run_h1_c = h1_corr.add_run("2. Detalles de Correcciones Críticas")
    run_h1_c.font.name = 'Arial'
    run_h1_c.font.size = Pt(15)
    run_h1_c.font.bold = True
    run_h1_c.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_corr.paragraph_format.space_before = Pt(18)
    h1_corr.paragraph_format.space_after = Pt(8)

    # Bullet 1: tred_equipment.id.number
    p_b1 = doc.add_paragraph(style='List Bullet')
    r_b1_title = p_b1.add_run("Número de Contenedor (tred_equipment.id.number): ")
    r_b1_title.bold = True
    p_b1.add_run(
        "En tu borrador estaba asignado a 'ParadaItem_Dyn.Z_PromocionCampana'. Esto sobrescribiría datos "
        "comerciales de campaña con la sigla física del contenedor (ej. MSKU1234567). Se ha corregido para "
        "apuntar a "
    )
    p_b1.add_run("Parada_Dyn.Z_Contenedor").bold = True
    p_b1.add_run(" (atributo custom en parada) o a la entidad nativa ")
    p_b1.add_run("UnidadContenedora.ReferenciaExterna").bold = True
    p_b1.add_run(".")

    # Bullet 2: tred_seal.number
    p_b2 = doc.add_paragraph(style='List Bullet')
    r_b2_title = p_b2.add_run("Asignación de Precintos (tred_seal.number): ")
    r_b2_title.bold = True
    p_b2.add_run(
        "Mapearlo a 'Viaje_Dyn.IdPrecinto' asume que todo el viaje terrestre lleva un único precinto. Si el "
        "viaje transporta múltiples contenedores (paradas), cada parada debe llevar su precinto independiente. "
        "Se recomienda utilizar "
    )
    p_b2.add_run("Parada_Dyn.Z_Precinto").bold = True
    p_b2.add_run(" o la relación nativa de recursos.")

    # Bullet 3: Empresa vs Transporte
    p_b3 = doc.add_paragraph(style='List Bullet')
    r_b3_title = p_b3.add_run("Concepto de Empresa y Transporte: ")
    r_b3_title.bold = True
    p_b3.add_run(
        "En UNIGIS, 'Empresa' representa a tu propia compañía (dador de carga/tenant). El transportista contratado "
        "debe guardarse en "
    )
    p_b3.add_run("Transporte.RazonSocial").bold = True
    p_b3.add_run(". Del mismo modo, el nombre del punto físico de carga/descarga (ej. terminal o depósito) es ")
    p_b3.add_run("Sucursal.Nombre").bold = True
    p_b3.add_run(" y no la razón social de la empresa.")

    # Bullet 4: Booking Number
    p_b4 = doc.add_paragraph(style='List Bullet')
    r_b4_title = p_b4.add_run("Booking Number (trcd_reference - BN): ")
    r_b4_title.bold = True
    p_b4.add_run(
        "Se completó el campo vacío mapeándolo a "
    )
    p_b4.add_run("Viaje.ReferenciaAdicional").bold = True
    p_b4.add_run(" o ")
    p_b4.add_run("Viaje.Z_BookingNumber").bold = True
    p_b4.add_run(", dado que la referencia externa principal del viaje está ocupada por el número único del mensaje XML.")

    # Bullet 5: Typos
    p_b5 = doc.add_paragraph(style='List Bullet')
    r_b5_title = p_b5.add_run("Errores Tipográficos y Casing: ")
    r_b5_title.bold = True
    p_b5.add_run(
        "Se corrigieron errores de tipeo y minúsculas: 'ReferenciaExteterna' a 'ReferenciaExterna', "
        "'Idcliente' a 'IdCliente' y 'TipoITem' a 'TipoItem'."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 3: Campos Adicionales
    h1_add = doc.add_paragraph()
    run_h1_a = h1_add.add_run("3. Campos Adicionales Recomendados")
    run_h1_a.font.name = 'Arial'
    run_h1_a.font.size = Pt(15)
    run_h1_a.font.bold = True
    run_h1_a.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_add.paragraph_format.space_before = Pt(18)
    h1_add.paragraph_format.space_after = Pt(8)

    p_add_intro = doc.add_paragraph()
    p_add_intro.add_run(
        "Para lograr un mapeo robusto del contenedor, sugerimos añadir al flujo XML los siguientes campos típicos "
        "de IFTMIN:"
    )
    p_add_intro.paragraph_format.space_after = Pt(6)

    p_a1 = doc.add_paragraph(style='List Bullet')
    p_a1.add_run("tred_equipment.size.and.type").bold = True
    p_a1.add_run(" (Tamaño/Tipo ISO, ej. 45G1, 20GP) ➔ Mapear a ")
    p_a1.add_run("Parada_Dyn.Z_TipoContenedor").bold = True
    p_a1.add_run(" o ")
    p_a1.add_run("TipoUnidadContenedora.ReferenciaExterna").bold = True
    p_a1.add_run(".")

    p_a2 = doc.add_paragraph(style='List Bullet')
    p_a2.add_run("tred_full.or.empty.indicator").bold = True
    p_a2.add_run(" (Indicador de Lleno/Vacío) ➔ Mapear a ")
    p_a2.add_run("Parada_Dyn.Z_EstadoEquipo").bold = True
    p_a2.add_run(" (valores: Full / Empty).")

    p_a3 = doc.add_paragraph(style='List Bullet')
    p_a3.add_run("tred_measurements").bold = True
    p_a3.add_run(" (con Calificador AET para peso neto) ➔ Mapear a ")
    p_a3.add_run("Parada_Dyn.Z_PesoNeto").bold = True
    p_a3.add_run(".")

    # Guardar documento en la ruta solicitada
    dest_dir = r"c:\Users\daniel.delamo\.gemini\antigravity\scratch\UniTask\docs\mapping"
    os.makedirs(dest_dir, exist_ok=True)
    
    file_path = os.path.join(dest_dir, "Mapeo_IFTMIN_UNIGIS.docx")
    doc.save(file_path)
    
    # También guardamos una copia en la raíz para fácil acceso
    root_path = r"c:\Users\daniel.delamo\.gemini\antigravity\scratch\UniTask\Mapeo_IFTMIN_UNIGIS.docx"
    doc.save(root_path)
    
    print(f"Documento DOCX generado exitosamente en: {file_path}")
    print(f"Copia generada exitosamente en la raíz: {root_path}")

if __name__ == "__main__":
    create_mapping_document()
