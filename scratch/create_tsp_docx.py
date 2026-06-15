import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_background(cell, hex_color):
    """Establece el color de fondo de una celda usando manipulación XML de Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_tsp_mapping_document():
    # Inicializar documento
    doc = docx.Document()
    
    # Configurar márgenes a 1 pulgada (2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilo de fuente global (Arial, 10.5pt, gris oscuro)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Título Principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title.add_run("Mapeo de Integración: Rendiciones, Cobros y Liquidaciones")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Azul Marino
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(4)

    # Subtítulo
    subtitle = doc.add_paragraph()
    run_sub = subtitle.add_run("Especificación de Interfaz de Salida (POST) al ERP - Proyecto TSP")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    subtitle.paragraph_format.space_after = Pt(24)

    # 1. INTRODUCCIÓN Y CONTEXTO
    h1_intro = doc.add_paragraph()
    run_h1_1 = h1_intro.add_run("1. Introducción y Contexto de la Interfaz")
    run_h1_1.font.name = 'Arial'
    run_h1_1.font.size = Pt(14)
    run_h1_1.font.bold = True
    run_h1_1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_intro.paragraph_format.space_before = Pt(12)
    h1_intro.paragraph_format.space_after = Pt(8)

    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "El presente documento especifica el mapeo de campos y el funcionamiento lógico de la interfaz POST "
        "utilizada para enviar los datos de "
    )
    p_intro.add_run("Cobros (Cobranzas) y Liquidaciones (Rendiciones de viaje)").bold = True
    p_intro.add_run(" desde el sistema de base de datos ")
    p_intro.add_run("UniGIS TMS (Proyecto TSP)").bold = True
    p_intro.add_run(
        " hacia el sistema ERP del cliente. El objetivo principal es la sincronización y conciliación financiera "
        "de los viajes una vez ejecutados."
    )
    p_intro.paragraph_format.space_after = Pt(12)

    p_flow = doc.add_paragraph()
    p_flow.add_run(
        "La base de datos del proyecto TSP es idéntica a la del entorno EUP (Europastry ES, bajo modelo de Shipper/Cargador), "
        "por lo que comparte su modelo físico de tablas base (dbo.Vehiculo, dbo.Conductor, dbo.Transporte, dbo.Cobros, dbo.Parada). "
        "La diferencia radica exclusivamente en los campos dinámicos custom (sufijos _Dyn) definidos para el negocio de TSP."
    )
    p_flow.paragraph_format.space_after = Pt(18)

    # 2. ARQUITECTURA DE COBROS Y RENDICIONES
    h1_arch = doc.add_paragraph()
    run_h1_2 = h1_arch.add_run("2. Arquitectura de Cobros y Rendiciones (Modelo de BD)")
    run_h1_2.font.name = 'Arial'
    run_h1_2.font.size = Pt(14)
    run_h1_2.font.bold = True
    run_h1_2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_arch.paragraph_format.space_before = Pt(18)
    h1_arch.paragraph_format.space_after = Pt(8)

    p_arch = doc.add_paragraph()
    p_arch.add_run(
        "En la base de datos transaccional del TMS (SQL Server), el circuito financiero del viaje se gestiona a través de las siguientes tablas:"
    )
    p_arch.paragraph_format.space_after = Pt(6)

    p_b1 = doc.add_paragraph(style='List Bullet')
    p_b1.add_run("dbo.Cobros: ").bold = True
    p_b1.add_run(
        "Almacena las transacciones de pago realizadas físicamente por los conductores a través del móvil (Efectivo, Tarjeta, TPV, Cheque, Pagaré). "
        "Contiene el importe cobrado, número de recibo, banco, fecha de vencimiento y foto del comprobante."
    )

    p_b2 = doc.add_paragraph(style='List Bullet')
    p_b2.add_run("dbo.ParadaCobro: ").bold = True
    p_b2.add_run("Tabla puente que asocia cada cobro realizado con la parada (dbo.Parada) del viaje correspondiente.")

    p_b3 = doc.add_paragraph(style='List Bullet')
    p_b3.add_run("dbo.Parada (TipoParada = 10 - Rendición): ").bold = True
    p_b3.add_run(
        "Al finalizar el viaje, un Stored Procedure (como dbo.Z_SP_GenerarRendicion_V2) evalúa los rechazos y mercancías sobrantes, "
        "generando una parada especial de tipo 10 (Rendición/Liquidación) con los bultos/pesos sobrantes a retornar al depósito de llegada."
    )

    p_b4 = doc.add_paragraph(style='List Bullet')
    p_b4.add_run("Relación con el Viaje: ").bold = True
    p_b4.add_run(
        "El viaje (dbo.Viaje) actúa como el eje agrupador de la ejecución. Relaciona el vehículo (dbo.Vehiculo), "
        "el conductor (dbo.Conductor) y la empresa transportista (dbo.Transporte). El ERP utiliza estas relaciones "
        "para liquidar las tarifas y consolidar los cobros."
    )

    p_b4.paragraph_format.space_after = Pt(18)

    # 3. TABLA DE MAPEO
    h1_map = doc.add_paragraph()
    run_h1_3 = h1_map.add_run("3. Tabla de Mapeo del JSON (POST al ERP)")
    run_h1_3.font.name = 'Arial'
    run_h1_3.font.size = Pt(14)
    run_h1_3.font.bold = True
    run_h1_3.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_map.paragraph_format.space_before = Pt(18)
    h1_map.paragraph_format.space_after = Pt(8)

    p_map_desc = doc.add_paragraph()
    p_map_desc.add_run(
        "A continuación se detalla cómo mapear cada propiedad del payload JSON con las tablas y columnas estándar y dinámicas "
        "de la base de datos de UniGIS / TSP (idéntica a EUP):"
    )
    p_map_desc.paragraph_format.space_after = Pt(12)

    headers = [
        "Nodo / Campo JSON",
        "Tabla de BD",
        "Columna de BD",
        "Tipo de Mapeo",
        "Explicación / Regla de Negocio"
    ]

    rows_data = [
        # Root Meta
        ("ApiKey", "N/A", "N/A", "Estático / Config", "Token de autorización para la API de comunicación con el ERP."),
        ("FechaDesde", "dbo.Viaje / dbo.Cobros", "Fecha / FechaHora", "Filtro de Fecha", "Inicio del rango de liquidaciones/cobros a procesar en el lote."),
        ("FechaHasta", "dbo.Viaje / dbo.Cobros", "Fecha / FechaHora", "Filtro de Fecha", "Fin del rango de liquidaciones/cobros a procesar en el lote."),
        ("CodigoSucursal", "dbo.Sucursal", "ReferenciaExterna", "Filtro / ID", "Código de la sucursal/depósito emisor de las operaciones."),
        ("EstadoGuias", "dbo.Guia / dbo.OrdenPedido", "IdEstado / Estado", "Filtro", "Código de estado de los documentos de entrega a incluir (entregados, cobrados)."),
        ("TipoGuia", "dbo.Guia", "IdTipoGuia (o .Tipo)", "Filtro", "Tipo de documento asociado al cobro (Factura, Remito, Albarán)."),
        ("CodigoOperacion", "dbo.Operacion", "ReferenciaExterna", "Filtro / ID", "Código de la operación comercial a la que pertenece la flota."),
        
        # Vehiculo
        ("Vehiculo.Dominio", "dbo.Vehiculo", "Dominio", "Estándar", "Matrícula/patente principal de la tractora (camión). Clave primaria de búsqueda."),
        ("Vehiculo.DominioSemi", "dbo.Vehiculo", "DominioSecundario", "Estándar", "Matrícula/patente del remolque/acoplado asignado."),
        ("Vehiculo.NroSerie", "dbo.Vehiculo", "NroSerie", "Estándar", "Número de chasis/serie del fabricante."),
        ("Vehiculo.Prestador", "dbo.Vehiculo", "Prestador", "Estándar", "Proveedor o prestador del servicio del vehículo."),
        ("Vehiculo.Flota", "dbo.Vehiculo", "Flota", "Estándar", "Clasificación de flota interna del dador de carga."),
        ("Vehiculo.Chasis", "dbo.Vehiculo", "Chasis", "Estándar", "Número de chasis del vehículo."),
        ("Vehiculo.Volumen", "dbo.Vehiculo", "Volumen", "Estándar", "Capacidad volumétrica en metros cúbicos."),
        ("Vehiculo.Peso", "dbo.Vehiculo", "Peso", "Estándar", "Capacidad de peso en kilogramos."),
        ("Vehiculo.Ciudad", "dbo.Vehiculo", "Ciudad", "Estándar", "Ciudad de base del vehículo."),
        ("Vehiculo.Marca", "dbo.Vehiculo", "Marca", "Estándar", "Marca del fabricante (ej: Scania, Volvo)."),
        ("Vehiculo.Modelo", "dbo.Vehiculo", "Modelo", "Estándar", "Modelo del vehículo."),
        ("Vehiculo.Combustible", "dbo.Vehiculo", "Combustible", "Estándar", "Tipo de combustible (Diésel, Eléctrico, etc.)."),
        ("Vehiculo.TipoVehiculo", "dbo.Vehiculo", "IdTipoVehiculo", "Relación / ID", "Clave foránea que resuelve en dbo.TipoVehiculo."),
        ("Vehiculo.Propietario", "dbo.Propietario", "Nombre / ReferenciaExterna", "Relación", "Titular legal del vehículo (dueño de la chapa)."),
        ("Vehiculo.Conductor", "dbo.Conductor", "Nombre / ReferenciaExterna", "Relación", "Conductor habitual asignado al vehículo."),
        ("Vehiculo.Transporte", "dbo.Transporte", "RazonSocial / Referencia", "Relación", "Transportista contratado responsable del equipo."),
        ("Vehiculo.ReferenciaExterna", "dbo.Vehiculo", "ReferenciaExterna", "Estándar", "ID integrador utilizado por el ERP para identificar el vehículo."),
        ("Vehiculo.Tara", "dbo.Vehiculo", "Tara", "Estándar", "Peso vacío (tara) del vehículo."),
        ("Vehiculo.IdEstadoVehiculo", "dbo.Vehiculo", "IdEstadoVehiculo", "Estándar", "ID del estado operacional en UniGIS."),
        
        # Vehiculo DYN / Opcionales
        ("Vehiculo.FechaFabricacion_opcional", "dbo.Vehiculo_Dyn", "FechaFabricacion", "Dinámico", "Fecha de fabricación del camión (para control de antigüedad en ERP)."),
        ("Vehiculo.IdEstado_opcional", "dbo.Vehiculo_Dyn", "IdEstado", "Dinámico", "Código de estado extendido en la tabla dinámica."),
        ("Vehiculo.IntegrarRNDC_Opcional", "dbo.Vehiculo_Dyn", "IntegrarRNDC", "Dinámico (Boolean)", "Indica si debe reportarse al Registro Nacional de Despacho de Carga (Argentina)."),
        ("Vehiculo.IntegrarCTE_Opcional", "dbo.Vehiculo_Dyn", "IntegrarCTE", "Dinámico (Boolean)", "Indica si debe reportarse al Conhecimento de Transporte (Brasil)."),
        ("Vehiculo.Varchar1 ... Varchar8", "dbo.Vehiculo", "Varchar1 ... Varchar8", "Estándar / Custom", "Campos de texto libre mapeables para datos del ERP."),
        ("Vehiculo.Int1, Int2", "dbo.Vehiculo", "Int1, Int2", "Estándar / Custom", "Campos enteros libres."),
        
        # Transporte
        ("Transporte.Referencia", "dbo.Transporte", "Referencia", "Estándar", "Código interno del transportista. Clave primaria lógica."),
        ("Transporte.Cuit", "dbo.Transporte", "Cuit", "Estándar", "Identificación fiscal de la empresa de transporte (CUIT/RUT)."),
        ("Transporte.RazonSocial", "dbo.Transporte", "RazonSocial", "Estándar", "Nombre legal del transportista."),
        ("Transporte.NombreFantasia", "dbo.Transporte", "NombreFantasia", "Estándar", "Nombre comercial."),
        ("Transporte.Email", "dbo.Transporte", "Email", "Estándar", "Contacto de facturación o liquidaciones."),
        ("Transporte.Direccion", "dbo.Transporte", "Direccion", "Estándar", "Domicilio fiscal de la empresa."),
        
        # Conductor
        ("Conductor.NroDocumento", "dbo.Conductor", "NroDocumento", "Estándar", "Número de documento (DNI/Cédula). Clave única del conductor."),
        ("Conductor.Login", "dbo.Conductor", "Login", "Estándar", "Usuario de la App Mobile que registró los cobros."),
        ("Conductor.Nombre", "dbo.Conductor", "Nombre", "Estándar", "Nombre del conductor."),
        ("Conductor.Apellido", "dbo.Conductor", "Apellido", "Estándar", "Apellido del conductor."),
        ("Conductor.Licencia", "dbo.Conductor", "Licencia", "Estándar", "Número de carnet/licencia de conducción."),
        ("Conductor.Vencimiento", "dbo.Conductor", "Vencimiento", "Estándar", "Fecha de vencimiento del carnet."),
        ("Conductor.ReferenciaExterna", "dbo.Conductor", "ReferenciaExterna", "Estándar", "ID integrador del conductor en el ERP.")
    ]

    # Crear tabla Word
    table = doc.add_table(rows=len(rows_data)+1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    col_widths = [Inches(1.8), Inches(1.2), Inches(1.3), Inches(1.0), Inches(1.2)]

    # Formatear Header
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "1F4E78") # Azul marino
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)

    # Rellenar filas
    for r_idx, row_val in enumerate(rows_data):
        row_cells = table.rows[r_idx+1].cells
        for c_idx, val in enumerate(row_val):
            row_cells[c_idx].text = val
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(8.5)
            
            # Highlight Tipo de Mapeo
            if c_idx == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    if val == "Dinámico":
                        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Rojo
                    elif val == "Relación":
                        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3) # Azul
                    elif val == "Estándar":
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Verde
                    else:
                        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Gris

    # Forzar anchos de columna
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. EXPLICACIÓN DE CAMPOS DINÁMICOS Y PERSONALIZACIÓN (DYN)
    h1_dyn = doc.add_paragraph()
    run_h1_4 = h1_dyn.add_run("4. Configuración de Campos Dinámicos (Tablas _Dyn)")
    run_h1_4.font.name = 'Arial'
    run_h1_4.font.size = Pt(14)
    run_h1_4.font.bold = True
    run_h1_4.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_dyn.paragraph_format.space_before = Pt(18)
    h1_dyn.paragraph_format.space_after = Pt(8)

    p_dyn = doc.add_paragraph()
    p_dyn.add_run(
        "Al igual que en EUP, la base de datos de TSP utiliza tablas de extensión con sufijo '_Dyn' para almacenar los datos "
        "que no caben en el estándar de UniGIS pero son críticos para el ERP. Las reglas de correspondencia son:"
    )
    p_dyn.paragraph_format.space_after = Pt(6)

    p_d1 = doc.add_paragraph(style='List Bullet')
    p_d1.add_run("dbo.Vehiculo_Dyn: ").bold = True
    p_d1.add_run(
        "Debe tener creadas las columnas para almacenar FechaFabricacion, IdEstadoVehiculo, Tara, y los flags "
        "de reportes fiscales (IntegrarRNDC, IntegrarCTE). Si el JSON contiene estos campos con valor, la interfaz los persiste en esta tabla relacionándolos por el Dominio o IdVehiculo."
    )

    p_d2 = doc.add_paragraph(style='List Bullet')
    p_d2.add_run("dbo.Conductor_Dyn: ").bold = True
    p_d2.add_run(
        "Almacena la fecha de nacimiento y los estados fiscales del conductor. Se relaciona mediante el IdConductor "
        "o NroDocumento."
    )

    p_d3 = doc.add_paragraph(style='List Bullet')
    p_d3.add_run("dbo.Transporte_Dyn: ").bold = True
    p_d3.add_run(
        "Almacena el estado administrativo en la publicación de ofertas de viajes (Tendering) y sus flags correspondientes."
    )

    # 5. INTEGRACIÓN CON COBROS Y LIQUIDACIONES EN EL ERP
    h1_erp = doc.add_paragraph()
    run_h1_5 = h1_erp.add_run("5. Sincronización y Conciliación con el ERP")
    run_h1_5.font.name = 'Arial'
    run_h1_5.font.size = Pt(14)
    run_h1_5.font.bold = True
    run_h1_5.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    h1_erp.paragraph_format.space_before = Pt(18)
    h1_erp.paragraph_format.space_after = Pt(8)

    p_erp_intro = doc.add_paragraph()
    p_erp_intro.add_run(
        "Cuando el ERP recibe el POST con este JSON, ejecuta los siguientes cruces en su módulo de finanzas y tesorería:"
    )
    p_erp_intro.paragraph_format.space_after = Pt(6)

    p_e1 = doc.add_paragraph(style='List Bullet')
    p_e1.add_run("Conciliación de Cobros: ").bold = True
    p_e1.add_run(
        "El ERP cruza los documentos de facturación (Guias/Pedidos) incluidos en la consulta con las transacciones registradas "
        "en dbo.Cobros para verificar el pago efectivo (Efectivo/TPV/Cheque). Si hay diferencias (Caso 2: menor al total, Caso 3: cobro en exceso), "
        "se generan las cuentas corrientes de cliente correspondientes."
    )

    p_e2 = doc.add_paragraph(style='List Bullet')
    p_e2.add_run("Liquidación de Transportistas: ").bold = True
    p_e2.add_run(
        "Utilizando el vehículo (Dominio), conductor (NroDocumento) y transportista (Referencia/Cuit), el ERP valida "
        "las tarifas acordadas en el contrato para calcular el pago del flete al transportista (dbo.Transporte), descontando "
        "posibles incidencias (rechazos/sobrantes) informadas en la parada de tipo 10 (Rendición)."
    )

    # Guardar documento
    dest_dir = r"c:\Users\daniel.delamo\.gemini\antigravity\scratch\UniTask\docs\mapping"
    os.makedirs(dest_dir, exist_ok=True)
    
    file_path = os.path.join(dest_dir, "Mapeo_Cobros_Liquidaciones_TSP.docx")
    doc.save(file_path)
    
    root_path = r"c:\Users\daniel.delamo\.gemini\antigravity\scratch\UniTask\Mapeo_Cobros_Liquidaciones_TSP.docx"
    doc.save(root_path)
    
    print(f"Documento de Mapeo generado en: {file_path}")
    print(f"Copia del Documento generada en la raíz: {root_path}")

if __name__ == "__main__":
    create_tsp_mapping_document()
