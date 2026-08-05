# -*- coding: utf-8 -*-
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    # Padding in dxa (1 pt = 20 dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_cell(cell, bg_color=None, bold=False, text_color=None, font_size=10, italic=False):
    if bg_color:
        set_cell_background(cell, bg_color)
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            if text_color:
                run.font.color.rgb = text_color

doc = Document()

# Page Setup
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Document Colors (UNIGIS Red: #E30613)
RED_HEX = "E30613"
RED_RGB = RGBColor(227, 6, 19)
GRAY_RGB = RGBColor(128, 128, 128)
BLACK_RGB = RGBColor(30, 30, 30)
WHITE_RGB = RGBColor(255, 255, 255)

# Title
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(4)
run_title = title_p.add_run("MATRIZ DE ESTADOS DE ENTIDADES — UNIGIS TMS")
run_title.font.name = 'Calibri'
run_title.font.size = Pt(20)
run_title.bold = True
run_title.font.color.rgb = RED_RGB

# Subtitle
sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(20)
run_sub = sub_p.add_run("Operación Internacional y 4PL (Transpais TSP)")
run_sub.font.name = 'Calibri'
run_sub.font.size = Pt(12)
run_sub.italic = True
run_sub.font.color.rgb = GRAY_RGB

# Intro Text
intro_p = doc.add_paragraph()
intro_p.paragraph_format.space_after = Pt(12)
run_intro = intro_p.add_run(
    "Este documento presenta la matriz completa de estados del flujo de transporte en UNIGIS TMS para Transpais. "
    "Se detalla la comparación entre los estados ciertos (validados en el DDS de Operación Internacional y 4PL) "
    "y los estados propuestos/imaginados que optimizan la gestión del sistema."
)
run_intro.font.name = 'Calibri'
run_intro.font.size = Pt(11)

# Section 1 Header
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(6)
run_h1 = h1.add_run("1. Matriz Resumen de Estados por Entidad")
run_h1.font.name = 'Calibri'
run_h1.font.size = Pt(14)
run_h1.bold = True
run_h1.font.color.rgb = RED_RGB

# Table 1: Entity States
table1 = doc.add_table(rows=5, cols=4)
add_table_borders(table1)

headers1 = ["Entidad", "Estados Ciertos (DDS Real)", "Estados Propuestos (TMS)", "Descripción / Lógica Operativa"]
row_data1 = [
    (
        "Pedidos\n(dbo.Pedido)",
        "INGRESADO, GRABADO, ERROR, CONFIRMADO, PLANIFICADO, LIQUIDADO",
        "EN RUTA, ENTREGADO, DEVUELTO, FACTURADO",
        "Los pedidos se tarifican en CONFIRMADO (forecast) y se recalculan al pasar a ENTREGADO con las cantidades reales."
    ),
    (
        "Viajes\n(dbo.Viaje)",
        "INACTIVO, ASIGNADO, CONFIRMADO, RECHAZADO, ACTIVO, EN EJECUCIÓN, RENDIDO, LIQUIDABLE, LIQUIDADO",
        "CANCELADO, REPROGRAMADO, AUDITADO",
        "La liquidación de compra se pre-calcula al activar el viaje y se cierra definitivamente cuando pasa a LIQUIDABLE."
    ),
    (
        "Paradas\n(dbo.Parada)",
        "VISITADO, Cargado, Cargado parcial, No cargado, Entregado, Entrega parcial, No entregado",
        "PENDIENTE, EN GEOCERCA, DEMORADA",
        "Registran la ejecución física del conductor desde la App UNIGIS X Deliveries o por geocerca automática."
    ),
    (
        "Órdenes\n(dbo.Orden)",
        "Se asocian directamente al pedido en planificación y a la parada en ejecución.",
        "INGRESADA, DESPACHADA, EN TRÁNSITO, CERRADA",
        "Entidad intermedia para gestionar tramos y sub-servicios en operativas complejas de Hook/Drop o 4PL."
    )
]

# Write headers for Table 1
hdr_cells = table1.rows[0].cells
for i, name in enumerate(headers1):
    hdr_cells[i].text = name
    format_cell(hdr_cells[i], bg_color=RED_HEX, bold=True, text_color=WHITE_RGB)

# Write body for Table 1
for idx, data in enumerate(row_data1):
    row_cells = table1.rows[idx + 1].cells
    bg = "F9F9F9" if idx % 2 == 1 else "FFFFFF"
    for j in range(4):
        row_cells[j].text = data[j]
        # Bold first column
        is_bold = (j == 0)
        t_color = RED_RGB if j == 0 else BLACK_RGB
        format_cell(row_cells[j], bg_color=bg, bold=is_bold, text_color=t_color)

# Add spacing
doc.add_paragraph().paragraph_format.space_before = Pt(12)

# Section 2 Header
h2 = doc.add_paragraph()
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)
run_h2 = h2.add_run("2. Detalle del Flujo de Ejecución de Paradas (DDS)")
run_h2.font.name = 'Calibri'
run_h2.font.size = Pt(14)
run_h2.bold = True
run_h2.font.color.rgb = RED_RGB

# Table 2: Stop Execution
table2 = doc.add_table(rows=17, cols=3)
add_table_borders(table2)

headers2 = ["Tipo de Parada (DDS)", "Estado Parada", "Motivo / Acción de Negocio"]
row_data2 = [
    ("CARGA EN REMITENTE", "Cargado en remitente", "Carga confirmada"),
    ("CARGA EN REMITENTE", "Cargado parcial en remitente", "Revisión TRANSPAIS — diferencia de bultos/palets"),
    ("CARGA EN REMITENTE", "No cargado en remitente", "Revisión TRANSPAIS — mercancía no disponible"),
    
    ("CARGA EN DEPÓSITO SALIDA", "Cargado en depósito salida", "Carga confirmada en COL"),
    ("CARGA EN DEPÓSITO SALIDA", "Cargado parcial en depósito salida", "Revisión TRANSPAIS"),
    ("CARGA EN DEPÓSITO SALIDA", "No cargado en depósito salida", "Revisión TRANSPAIS"),
    
    ("CARGA EN DEPÓSITO LLEGADA", "Cargado en depósito llegada", "Carga confirmada en depósito destino"),
    ("CARGA EN DEPÓSITO LLEGADA", "Cargado parcial en depósito llegada", "Revisión TRANSPAIS"),
    ("CARGA EN DEPÓSITO LLEGADA", "No cargado en depósito llegada", "Revisión TRANSPAIS"),
    
    ("ENTREGA EN CONSIGNATARIO", "Entregado en consignatario", "Entrega confirmada al cliente final"),
    ("ENTREGA EN CONSIGNATARIO", "Entrega parcial en consignatario", "Revisión TRANSPAIS — se crea parada devolución"),
    ("ENTREGA EN CONSIGNATARIO", "No entregado en consignatario", "Revisión TRANSPAIS — se crea parada devolución"),
    
    ("ENTREGA EN DEPÓSITO SALIDA", "Entregado en depósito salida", "Confirmado"),
    ("ENTREGA EN DEPÓSITO SALIDA", "Entrega parcial / No entregado", "Revisión TRANSPAIS"),
    
    ("ENTREGA EN DEPÓSITO LLEGADA", "Entregado en depósito llegada", "Confirmado"),
    ("ENTREGA EN DEPÓSITO LLEGADA", "Entrega parcial / No entregado", "Revisión TRANSPAIS")
]

# Write headers for Table 2
hdr_cells2 = table2.rows[0].cells
for i, name in enumerate(headers2):
    hdr_cells2[i].text = name
    format_cell(hdr_cells2[i], bg_color=RED_HEX, bold=True, text_color=WHITE_RGB)

# Write body for Table 2
last_type = ""
type_color_toggle = True
for idx, data in enumerate(row_data2):
    row_cells = table2.rows[idx + 1].cells
    current_type = data[0]
    
    # Visual grouping by Stop Type
    if current_type != last_type:
        type_color_toggle = not type_color_toggle
        last_type = current_type
    
    bg = "FFF5F5" if type_color_toggle else "FFFFFF"
    
    for j in range(3):
        row_cells[j].text = data[j]
        is_bold = (j == 0)
        t_color = RED_RGB if j == 0 else BLACK_RGB
        format_cell(row_cells[j], bg_color=bg, bold=is_bold, text_color=t_color)

# Save Document
doc.save("scratch/estados_entidades_unigis.docx")
print("Successfully created scratch/estados_entidades_unigis.docx")
